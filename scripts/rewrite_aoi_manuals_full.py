from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


DESKTOP = Path(r"c:\Users\Neizy\Desktop")

FILES = {
    "sysprog": DESKTOP / "АОИ.01.32 01 — Руководство системного программиста (АОИ-Web).docx",
    "programmer": DESKTOP / "АОИ.01.33 01 — Руководство программиста (АОИ-Web).docx",
    "operator": DESKTOP / "АОИ.01.34 01 — Руководство оператора (АОИ-Web).docx",
    "manager": DESKTOP / "АОИ.01.35 01 — Руководство руководителя (АОИ-Web).docx",
}


def _setup_doc(doc: Document, title: str, code: str) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    p = doc.add_paragraph("АОИ-WEB — ПРОГРАММНАЯ ЧАСТЬ ПАК АВТОМАТИЧЕСКОЙ ОПТИЧЕСКОЙ ИНСПЕКЦИИ")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True

    p = doc.add_paragraph(title)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True

    p = doc.add_paragraph(code)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()


def _h(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def _p(doc: Document, text: str) -> None:
    doc.add_paragraph(text)


def _bullets(doc: Document, items: list[str]) -> None:
    for i in items:
        doc.add_paragraph(i, style="List Bullet")


def build_sysprog() -> Document:
    d = Document()
    _setup_doc(d, "Руководство системного программиста", "АОИ.01.32 01")

    _h(d, "Аннотация")
    _p(
        d,
        "Документ содержит сведения по системной настройке, развёртыванию, проверке и сопровождению "
        "программного комплекса АОИ-Web. Комплекс предназначен для автоматической оптической инспекции "
        "печатных узлов с поддержкой web-интерфейса, нейросетевой детекции YOLOv8, опционального "
        "выравнивания по эталону (Golden Board) и управления подсветкой через WLED.",
    )

    _h(d, "1. Общие сведения о программе")
    _p(
        d,
        "АОИ-Web — серверное приложение на FastAPI с клиентской SPA-частью. Комплекс обеспечивает "
        "инспекцию плат, хранение протоколов, отчётность, управление пользователями и настройками оборудования.",
    )
    _bullets(
        d,
        [
            "Назначение: автоматизация визуального контроля печатных узлов.",
            "Технологии: Python 3.10+, FastAPI, SQLAlchemy, Alembic, PyTorch, Ultralytics YOLOv8, OpenCV.",
            "Роли доступа: оператор, руководитель, администратор.",
            "Интеграции: WLED JSON API (/json/info, /json/state), мобильная страница /phone.",
        ],
    )

    _h(d, "2. Структура программы")
    _bullets(
        d,
        [
            "app/main.py — точка входа FastAPI, подключение роутеров, статика.",
            "app/api/* — API: auth, users, inspections, datasets, pipeline, stats.",
            "app/services/* — детектор, preprocessing, alignment, gateway подсветки.",
            "app/models/* и alembic/* — модели и миграции БД.",
            "storage/, models/, logs/ — рабочие каталоги данных.",
        ],
    )

    _h(d, "3. Настройка программы")
    _bullets(
        d,
        [
            "Создать .env на основе .env.example и задать SECRET_KEY.",
            "Инициализировать БД: python -m scripts.init_db.",
            "Запуск dev: uvicorn app.main:app --reload --host 0.0.0.0 --port 8000.",
            "Portable: запуск AOI-Web-Portable-HTTPS.exe.",
            "Для доступа телефона задать PUBLIC_BASE_URL (LAN IP).",
            "Настройка WLED через админ-панель: адрес, /json/info, /json/state, segment/transition.",
        ],
    )

    _h(d, "4. Проверка программы")
    _bullets(
        d,
        [
            "Проверить доступность /, /docs, /redoc.",
            "Выполнить тестовую инспекцию с загрузкой изображения.",
            "Проверить, что создаются записи в БД и файлы в storage.",
            "Для WLED: hardware/probe, lighting/control, hardware/admin/diagnostics.",
            "Проверить лог запуска: Application startup complete.",
        ],
    )

    _h(d, "5. Сообщения системному программисту")
    _bullets(
        d,
        [
            "Application startup complete — запуск успешен.",
            "Порт недоступен — изменить AOI_WEB_PORT или освободить порт.",
            "Ошибка импорта app.main — проверить целостность поставки и зависимости.",
            "Ошибка связи с WLED — проверить IP, сеть, доступность /json/info.",
        ],
    )

    return d


def build_programmer() -> Document:
    d = Document()
    _setup_doc(d, "Руководство программиста", "АОИ.01.33 01")

    _h(d, "Аннотация")
    _p(
        d,
        "Документ предназначен для программиста, выполняющего разработку и доработку АОИ-Web: "
        "API, сервисов детекции, предобработки, выравнивания, отчётности и интеграции внешнего оборудования.",
    )

    _h(d, "1. Назначение и условия применения программы")
    _bullets(
        d,
        [
            "Комплекс применяется в составе ПАК АОИ для контроля качества печатных узлов.",
            "Серверный режим: FastAPI/Uvicorn, SQLite или PostgreSQL.",
            "ML-часть: YOLOv8, переключаемые датасеты и веса.",
        ],
    )

    _h(d, "2. Характеристика программы")
    _bullets(
        d,
        [
            "Клиент-серверная архитектура; SPA интерфейс + REST API.",
            "Ролевой доступ на JWT: operator/manager/admin.",
            "Поддержка Golden Board (ECC alignment), WLED, отчётов PDF/CSV.",
            "Логирование в logs/aoi.log.",
        ],
    )

    _h(d, "3. Обращение к программе")
    _bullets(
        d,
        [
            "Установка зависимостей: pip install -r requirements.txt.",
            "Инициализация БД: python -m scripts.init_db.",
            "Запуск: uvicorn app.main:app --host 0.0.0.0 --port 8000.",
            "API-контроль: /docs и /openapi.json.",
            "Portable-сборка: scripts/build_portable_https.ps1.",
        ],
    )

    _h(d, "4. Входные и выходные данные")
    _bullets(
        d,
        [
            "Вход: изображения плат, настройки .env, JSON-запросы API, конфиг WLED.",
            "Выход: протоколы инспекций, данные дефектов, экспортные отчёты, журналы.",
        ],
    )

    _h(d, "5. Сообщения")
    _bullets(
        d,
        [
            "Application startup complete — успешный запуск.",
            "Недоступен порт — настройка AOI_WEB_PORT/AOI_WEB_PORT_FALLBACK.",
            "Таймаут WLED — проверка адреса/сети/контроллера.",
        ],
    )

    return d


def build_operator() -> Document:
    d = Document()
    _setup_doc(d, "Руководство оператора", "АОИ.01.34 01")

    _h(d, "Аннотация")
    _p(
        d,
        "Документ описывает порядок работы оператора в АОИ-Web: вход в систему, запуск инспекции, "
        "получение изображения (в том числе через телефон), просмотр результатов и фиксация данных.",
    )

    _h(d, "1. Назначение программы")
    _p(d, "Оператор выполняет инспекцию печатных узлов и фиксирует результаты проверки в системе.")

    _h(d, "2. Условия выполнения программы")
    _bullets(
        d,
        [
            "Доступ к серверу АОИ-Web по локальной сети.",
            "Браузер Chrome/Edge/Firefox.",
            "Учётная запись оператора.",
            "При съёмке через телефон — разрешение на доступ к камере.",
        ],
    )

    _h(d, "3. Выполнение программы")
    _bullets(
        d,
        [
            "Войти в систему под оператором.",
            "Создать инспекцию в разделе «Инспекция».",
            "Сделать снимок через /phone или загрузить файл.",
            "Запустить анализ и дождаться результата.",
            "Проверить список найденных дефектов и сохранить протокол.",
        ],
    )

    _h(d, "4. Сообщения оператору")
    _bullets(
        d,
        [
            "Неверный логин/пароль — повторить вход.",
            "Камера недоступна — выдать разрешение в браузере.",
            "Ошибка загрузки изображения — проверить формат/размер файла.",
        ],
    )

    return d


def build_manager() -> Document:
    d = Document()
    _setup_doc(d, "Руководство руководителя", "АОИ.01.35 01")

    _h(d, "Аннотация")
    _p(
        d,
        "Документ описывает работу руководителя в АОИ-Web: просмотр сводного журнала инспекций, "
        "анализ статистики, контроль датасетов и устройств, принятие решений по повторным проверкам.",
    )

    _h(d, "1. Назначение")
    _p(
        d,
        "Руководитель использует АОИ-Web для контроля качества, анализа динамики дефектов "
        "и координации работы операторов.",
    )

    _h(d, "2. Доступные функции")
    _bullets(
        d,
        [
            "Просмотр общего журнала инспекций.",
            "Просмотр статистики дефектов.",
            "Работа с датасетами (в пределах прав role=manager).",
            "Контроль устройств и параметров съёмки/инспекции.",
            "Экспорт отчётных данных.",
        ],
    )

    _h(d, "3. Порядок работы")
    _bullets(
        d,
        [
            "Войти в систему под ролью руководителя.",
            "Проверить разделы «Журнал», «Статистика», «Датасеты», «Устройства».",
            "Сформировать замечания по инспекциям и назначить повторные проверки при необходимости.",
        ],
    )

    _h(d, "4. Сообщения руководителю")
    _bullets(
        d,
        [
            "Недостаточно прав — обратиться к администратору.",
            "Ошибка загрузки/доступа к данным — проверить связь с сервером и повторить операцию.",
        ],
    )

    return d


def main() -> None:
    docs = {
        "sysprog": build_sysprog(),
        "programmer": build_programmer(),
        "operator": build_operator(),
        "manager": build_manager(),
    }
    for key, doc in docs.items():
        out = FILES[key]
        doc.save(out)
        print(f"saved: {out}")


if __name__ == "__main__":
    main()

