"""Generate AOI-Web system programmer manual (.docx) for diploma docs."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT


def add_para(doc: Document, text: str, bold: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(12)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def build() -> Document:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    # Title page block (simplified)
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("УТВЕРЖДЕНО\nАОИ.01.32 01-ЛУ")
    r.bold = True

    doc.add_paragraph()
    t2 = doc.add_paragraph()
    t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = t2.add_run(
        "АОИ-WEB — ПРОГРАММНАЯ ЧАСТЬ ПАК\n"
        "АВТОМАТИЧЕСКОЙ ОПТИЧЕСКОЙ ИНСПЕКЦИИ\n\n"
        "Руководство системного программиста\n"
        "АОИ.01.32 01"
    )
    r2.bold = True

    doc.add_page_break()

    add_heading(doc, "АННОТАЦИЯ", 1)
    for block in [
        "В данном программном документе приведено руководство системного программиста "
        "по настройке и использованию программного комплекса «АОИ-Web» (обозначение: АОИ.01), "
        "предназначенного для автоматической оптической инспекции печатных узлов.",
        "Комплекс обеспечивает захват изображения через веб-интерфейс (в том числе с камеры "
        "мобильного устройства), предварительную обработку кадра, детекцию объектов и дефектов "
        "с применением нейросетевой модели (Ultralytics YOLOv8), опциональное выравнивание кадра "
        "по эталону (Golden Board), формирование протоколов инспекции и ведение журнала. "
        "Также комплекс поддерживает управление подсветкой стенда через контроллер WLED "
        "по локальной сети (JSON API).",
        "В разделе «Общие сведения о программе» указаны назначение и функции комплекса, "
        "сведения о технических и программных средствах, а также требования к персоналу.",
        "В разделе «Структура программы» приведены сведения о структуре, составных частях "
        "и связях с другими программами и устройствами.",
        "В разделе «Настройка программы» описаны действия по настройке на условия конкретного "
        "применения (окружение, сеть, интеграция WLED, выбор режима детектора).",
        "В разделе «Проверка программы» приведены способы проверки работоспособности.",
        "В разделе «Сообщения системному программисту» указаны тексты сообщений и действия "
        "по ним.",
        "Оформление документа выполнено по требованиям ЕСПД (ГОСТ 19.101-77, 19.103-77, "
        "19.104-78, 19.105-78, 19.106-78, 19.503-79, 19.604-78).",
    ]:
        add_para(doc, block)

    doc.add_page_break()
    add_heading(doc, "СОДЕРЖАНИЕ", 1)
    toc = [
        "1. Общие сведения о программе",
        "2. Структура программы",
        "3. Настройка программы",
        "4. Проверка программы",
        "5. Сообщения системному программисту",
    ]
    for line in toc:
        add_para(doc, line)

    doc.add_page_break()
    add_heading(doc, "1. ОБЩИЕ СВЕДЕНИЯ О ПРОГРАММЕ", 1)

    add_heading(doc, "1.1. Назначение программы", 2)
    add_para(
        doc,
        "Программный комплекс «АОИ-Web» предназначен для автоматизации оптического "
        "контроля печатных узлов в составе ПАК автоматической оптической инспекции.",
    )
    add_bullets(
        doc,
        [
            "приём изображений (включая съёмку через страницу /phone);",
            "предобработка кадра и детекция дефектов (YOLOv8);",
            "выравнивание по эталону Golden Board (опционально);",
            "формирование протокола инспекции;",
            "управление подсветкой WLED (опционально).",
        ],
    )

    add_heading(doc, "1.2. Функции программы", 2)
    add_bullets(
        doc,
        [
            "веб-интерфейс и HTTP API (FastAPI);",
            "аутентификация и роли (JWT);",
            "инспекции, хранение в БД и storage;",
            "детекция YOLOv8, выбор активного датасета;",
            "отчёты и журналирование;",
            "интеграция WLED (GET /json/info, POST /json/state).",
        ],
    )

    add_heading(doc, "1.3. Минимальный состав технических средств", 2)
    add_bullets(
        doc,
        [
            "ПК x64, ОЗУ ≥ 8 Гбайт (рекомендуется);",
            "диск ≥ 5 Гбайт свободного места;",
            "сеть LAN; устройство захвата (смартфон/камера).",
        ],
    )

    add_heading(doc, "1.4. Минимальный состав программных средств", 2)
    add_bullets(
        doc,
        [
            "Windows 10/11 (portable) или Python 3.10+ (из исходников);",
            "браузер Chrome/Edge/Firefox;",
            "WLED с HTTP JSON API (при использовании подсветки).",
        ],
    )

    add_heading(doc, "1.5. Требования к персоналу (системному программисту)", 2)
    add_para(
        doc,
        "Системный программист должен уметь устанавливать и сопровождать ПО под Windows, "
        "настраивать сеть и .env, анализировать логи, обеспечивать работу WLED и доступ "
        "с мобильных устройств.",
    )

    doc.add_page_break()
    add_heading(doc, "2. СТРУКТУРА ПРОГРАММЫ", 1)

    add_heading(doc, "2.1. Сведения о структуре программы", 2)
    add_para(
        doc,
        "Клиент–серверная архитектура: FastAPI/Uvicorn, SPA, БД, storage, сервисы "
        "детекции и выравнивания, шлюз WLED.",
    )

    add_heading(doc, "2.2. Сведения о составных частях программы", 2)
    add_bullets(
        doc,
        [
            "app/main.py — точка входа;",
            "app/api/* — маршруты API;",
            "app/services/* — детектор, preprocessing, alignment, gateway;",
            "app/models/*, alembic/* — БД;",
            "app/static/* — интерфейс.",
        ],
    )

    add_heading(doc, "2.3. Сведения о связях между составными частями программы", 2)
    add_para(
        doc,
        "API вызывает сервисы; сервисы используют настройки, БД и storage; детектор "
        "синхронизируется с активным датасетом из БД.",
    )

    add_heading(doc, "2.4. Сведения о связях с другими программами", 2)
    add_bullets(
        doc,
        [
            "веб-браузеры (HTTP/HTTPS);",
            "WLED (JSON API);",
            "средства ОС (порты, firewall).",
        ],
    )

    doc.add_page_break()
    add_heading(doc, "3. НАСТРОЙКА ПРОГРАММЫ", 1)

    add_heading(doc, "3.1. Настройка на состав технических средств", 2)
    add_para(
        doc,
        "Обеспечить сетевой доступ к ПК, при съёмке с телефона — корректный "
        "PUBLIC_BASE_URL; для WLED — доступность IP в LAN.",
    )

    add_heading(doc, "3.2. Настройка на состав программных средств", 2)
    add_para(doc, "Portable: запуск AOI-Web-Portable-HTTPS.exe из каталога поставки.")
    add_bullets(
        doc,
        [
            "AOI_WEB_PORT — порт;",
            "AOI_WEB_PORT_FALLBACK=1 — авто-смена занятого порта;",
            "PUBLIC_BASE_URL — LAN-адрес для /phone;",
            "настройка WLED в админ-панели.",
        ],
    )
    add_para(
        doc,
        "Из исходников: .env, pip install -r requirements.txt, "
        "uvicorn app.main:app --host 0.0.0.0 --port 8000.",
    )

    doc.add_page_break()
    add_heading(doc, "4. ПРОВЕРКА ПРОГРАММЫ", 1)

    add_heading(doc, "4.1. Описание способов проверки", 2)
    add_bullets(
        doc,
        [
            "запуск и доступность / и /docs;",
            "контрольная инспекция (загрузка изображения);",
            "проверка WLED (связь и POST /json/state).",
        ],
    )

    add_heading(doc, "4.2. Методы прогона", 2)
    add_heading(doc, "4.2.1. Проверка работоспособности программы", 3)
    add_bullets(
        doc,
        [
            "запустить комплекс;",
            "открыть интерфейс и /docs;",
            "выполнить тестовую инспекцию;",
            "при наличии WLED — «Проверить связь».",
        ],
    )

    add_heading(doc, "4.2.2. Проверка на сообщение об ошибке", 3)
    add_bullets(
        doc,
        [
            "занять порт 8000 или задать занятый AOI_WEB_PORT;",
            "указать неверный адрес WLED и выполнить проверку связи.",
        ],
    )

    doc.add_page_break()
    add_heading(doc, "5. СООБЩЕНИЯ СИСТЕМНОМУ ПРОГРАММИСТУ", 1)

    messages = [
        (
            "Application startup complete.",
            "Успешный запуск. Перейти к проверке интерфейса.",
        ),
        (
            "Порт N недоступен... AOI_WEB_PORT / AOI_WEB_PORT_FALLBACK",
            "Порт занят. Освободить порт или включить fallback.",
        ),
        (
            "Ошибка импорта app.main",
            "Повреждена поставка или отсутствуют зависимости. Проверить _internal.",
        ),
        (
            "Ошибки HTTP к WLED (/json/info, /json/state)",
            "Проверить IP WLED, сеть и firewall.",
        ),
    ]
    for i, (msg, action) in enumerate(messages, 1):
        add_para(doc, f"{i}) «{msg}»", bold=True)
        add_para(doc, f"ДЕЙСТВИЯ ПРОГРАММИСТА: {action}")

    doc.add_page_break()
    add_heading(doc, "Лист регистрации изменений", 1)
    add_para(doc, "(заполняется при внесении изменений в документ)")

    return doc


def main() -> None:
    out = Path(r"c:\Users\Neizy\Desktop\АОИ.01.32 01 — Руководство системного программиста (АОИ-Web).docx")
    build().save(str(out))
    print(f"Saved: {out}")


if __name__ == "__main__":
    main()
